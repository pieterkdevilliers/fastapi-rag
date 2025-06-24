import os
import requests
import io
import uuid
import shutil
import gc
import numpy as np
import boto3
import openai
import asyncio
from docx import Document as DocxDocument
import PyPDF2
import pypandoc
import markdown
import chromadb
from langchain_chroma import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import OpenAIEmbeddings
from sqlalchemy.sql import select
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.asyncio import AsyncSession
from dotenv import load_dotenv
from file_management.models import SourceFile
from db import async_engine
from botocore.exceptions import ClientError
from chromadb.api.types import EmbeddingFunction

load_dotenv()

ENVIRONMENT = os.environ.get('ENVIRONMENT')

# Chroma API endpoint and credentials
CHROMA_ENDPOINT = os.environ.get('CHROMA_ENDPOINT')
CHROMA_SERVER_AUTHN_CREDENTIALS = os.environ.get('CHROMA_SERVER_AUTHN_CREDENTIALS')

headers = {
    'X-Chroma-Token': CHROMA_SERVER_AUTHN_CREDENTIALS,
    'Content-Type': 'application/json'
}

embedding_function = OpenAIEmbeddings()
# sample_text = "Sample text to check embedding size."
# embedding = embedding_function.embed_documents([sample_text])
# print(f"Embedding size: {len(embedding[0])}")

class ChromaEmbeddingFunction(EmbeddingFunction):
    def __init__(self):
        self.embedding_function = OpenAIEmbeddings()

    def __call__(self, input):
        # Ensure that the input is a list of strings
        if not isinstance(input, list):
            input = [input]
        return self.embedding_function.embed_documents(input)
    
    def get_dimension(self):
        return self.embedding_function.get_dimension()

# Initialize the S3 client
s3 = boto3.client('s3')
BUCKET_NAME = os.environ.get('AWS_STORAGE_BUCKET_NAME')

# Create a sessionmaker for the async session
AsyncSessionLocal = sessionmaker(
    bind=async_engine,
    class_=AsyncSession,
    expire_on_commit=False
)

async def get_async_session() -> AsyncSession:
    """
    Get Async Session
    """
    session = AsyncSessionLocal()
    try:
        yield session
    finally:
        await session.close()

openai.api_key = os.getenv('OPENAI_API_KEY')
if not openai.api_key:
    raise ValueError("OPENAI_API_KEY environment variable not set")

# BATCH_SIZE = 5461
BATCH_SIZE = 100

async def generate_chroma_db(account_unique_id, replace=False):
    """
    Generate a data store from the documents.
    """
    print(f"generate_chroma_db called for {account_unique_id}.")

    async for session in get_async_session():
        await generate_data_store(account_unique_id, replace, session)
    
    response = {"response": "success"}
    return response

async def generate_data_store(account_unique_id: str, replace: bool, session: AsyncSession):
    """
    Generate a data store from the documents uploaded to S3.
    """
    print(f"Generating collection store for account {account_unique_id}.")

    # Determine the Chroma path based on the environment
    if ENVIRONMENT == 'development':
        chroma_path = f"./chroma/{account_unique_id}"
    else:
        chroma_path = (f'{CHROMA_ENDPOINT}')
        print(f"Chroma path: {chroma_path}")

    # Fetch the document URLs from the database instead of reading from local file_path
    documents = await load_documents_from_s3(account_unique_id, replace, session)

    # Split the text into chunks
    chunks = await split_text(documents)

    # Check for empty chunks before initializing the database
    if chunks:
        await save_to_chroma_in_batches(chunks, chroma_path, replace, account_unique_id)
    else:
        print(f"No chunks generated for account {account_unique_id}. Skipping Chroma DB creation.")


async def save_to_chroma_in_batches(chunks: list[Document], chroma_path: str, replace: bool, account_unique_id: str):
    """
    Save the chunks to the Chroma database in batches.
    """
    print(f"Saving chunks to Chroma DB at {chroma_path} save_to_chroma_in_batches.")
    embedding_function = ChromaEmbeddingFunction()

    # Check if we are using a local or remote ChromaDB
    if ENVIRONMENT == 'development':
        # Handle local ChromaDB setup
        if not os.path.exists(chroma_path):
            await create_chroma_db_dir(chroma_path)
            print(f"Directory created: {chroma_path}")
        else:
            print(f"Directory already exists: {chroma_path}")

        db = None
        if replace:
            print(f"Creating new Chroma DB at {chroma_path}.")
            db = await create_new_chroma_db(chroma_path, embedding_function)
        else:
            print(f"Loading existing Chroma DB at {chroma_path}.")
            db = Chroma(persist_directory=chroma_path, embedding_function=embedding_function)

        if db is None:
            print("Error initializing Chroma DB.")
            return {"response": "error", "message": "Failed to initialize Chroma DB."}

        await save_chunks_to_local_db(chunks, db, chroma_path)
    else:
        # Handle remote ChromaDB setup using HTTP API calls
        print("Using remote ChromaDB setup.")
        data = {
            "database": "default_database",
            "name": (f"collection-{account_unique_id}"),
            "tenant": "default_tenant",
            }
        print("data:", data)
        
        client = chromadb.HttpClient(host='https://fastapi-rag-chroma.onrender.com', port=8000, headers=headers)
        print("client:", client)
        try:
            collection = client.get_collection(name=f"collection-{account_unique_id}", embedding_function=embedding_function)
            print("collection:", collection.id)
            collection_id = collection.id
        except Exception as e:
            if "does not exist" in str(e).lower():
                embedding_function = ChromaEmbeddingFunction()  # Should return embeddings of dimension 1536

                # Create the Chroma collection with the correct embedding function
                data = {
                    "name": f"collection-{account_unique_id}",
                    "tenant": "default_tenant",
                    "database": "default_database",
                    # "dimension": 1536,  # Ensure the correct dimension is set
                }

                collection = client.create_collection(name=f"collection-{account_unique_id}", embedding_function=embedding_function)
                print("collection from create:", collection.id)
                collection_id = collection.id
                print(f"Collection ID: {collection_id}")
            print(f"Error retrieving collection: {str(e)}")
            return {"response": "error", "message": str(e)}
    
        collection_id = collection.id
        await save_chunks_to_remote_db(chunks, chroma_path, replace, headers=headers, json=data, collection_id=collection_id)


async def save_chunks_to_local_db(chunks: list[Document], db: Chroma, chroma_path: str):
    """
    Save the chunks to the local Chroma database.
    """
    async for chunk_batch in batch(chunks):
        try:
            db.add_documents(chunk_batch)
            print(f"Saved {len(chunk_batch)} chunks to local ChromaDB at {chroma_path}.")
        except Exception as e:
            print(f"Error saving chunks to local Chroma DB: {e}")
            return {"response": "error", "message": str(e)}


async def save_chunks_to_remote_db(chunks: list[Document], chroma_path: str, replace: bool, headers: dict, json: dict, collection_id: str):
    """
    Save the chunks to the remote Chroma database using API calls.
    """
    print(f"Saving chunks to remote Chroma DB at {chroma_path}.")
    embeddings_model = OpenAIEmbeddings()
    # embeddings_model = ChromaEmbeddingFunction()
    expected_dimension = 1536  # The dimension defined in your collection
    
    async for chunk_batch in batch(chunks):
        try:
            # Ensure each chunk has a unique ID
            for chunk in chunk_batch:
                if chunk.id is None:
                    chunk.id = str(uuid.uuid4())

            # Generate embeddings for each chunk
            embedding_vectors = await asyncio.gather(
                *[embeddings_model.aembed_documents(chunk.page_content) for chunk in chunk_batch]
            )

            # Log the dimensions of the embeddings
            for vec in embedding_vectors:
                print(f"Initial embedding shape: {np.array(vec).shape}")

            # Resize embeddings if necessary
            resized_embeddings = []
            for vec in embedding_vectors:
                # Convert to numpy array if it is a list
                vec = np.array(vec)

                # Flatten the array to ensure it's one-dimensional
                if vec.shape[0] != expected_dimension:
                    print(f"Mismatch! Expected dimension: {expected_dimension}, but got: {vec.shape[0]}")
                
                # Resize or trim the embedding to match the expected dimension
                resized = vec.flatten()[:expected_dimension]  # Take the first 'expected_dimension' elements

                resized_embeddings.append(resized)
                print(f"Resized embedding shape: {resized.shape}")

            # Preparing the payload for the API
            payload = {
                "ids": [chunk.id for chunk in chunk_batch],
                "documents": [chunk.page_content for chunk in chunk_batch],
                "metadatas": [chunk.metadata for chunk in chunk_batch],
                "embeddings": [vec.tolist() for vec in resized_embeddings],  # Convert to list for JSON serialization
                "uris": []  # Optional: add URIs if applicable
            }

            response = requests.post(
                f"{chroma_path}/collections/{collection_id}/add", 
                json=payload, 
                headers=headers
            )
            if response.status_code != 201:
                print(f"Error saving chunks to remote Chroma DB: {response.text}")
                return {"response": "error", "message": response.text}
            else:
                print(f"Saved {len(chunk_batch)} chunks to remote ChromaDB at {collection_id}.")
                gc.collect()
        except Exception as e:
            print(f"Error saving chunks to remote Chroma DB: {e}")
            return {"response": "error", "message": str(e)}


def resize_embedding(embedding, expected_size):
    # Convert to numpy array if it is a list
    if isinstance(embedding, list):
        embedding = np.array(embedding)

    current_size = embedding.shape[0]
    
    if current_size < expected_size:
        # Pad with zeros if too small
        resized = np.pad(embedding, (0, expected_size - current_size), 'constant')
    elif current_size > expected_size:
        # Trim if too large
        resized = embedding[:expected_size]
    else:
        resized = embedding
    
    # Flatten to ensure it's a 1D array
    return resized.flatten()
        

async def create_chroma_db_dir(chroma_path: str):
    """
    Create Chroma DB directory locally.
    """
    os.makedirs(chroma_path, exist_ok=True)
    os.chmod(chroma_path, 0o775)
    print(f"Successfully created new directory: {chroma_path}.")
    return chroma_path

async def create_new_chroma_db(chroma_path: str, embeddings):
    """
    Create a new Chroma database locally.
    """
    try:
        db = Chroma(persist_directory=chroma_path, embedding_function=embeddings)
        print(f"Chroma DB initialized successfully at {chroma_path}.")
        return db
    except Exception as e:
        print(f"Error initializing Chroma DB: {e}")
        return None


async def load_documents_from_s3(account_unique_id: str, replace: bool, session: AsyncSession):
    """
    Load documents from S3 based on a database query.
    """
        # Query database for files to process
    if replace:
        statement = select(SourceFile).filter(
            SourceFile.account_unique_id == account_unique_id,
            SourceFile.included_in_source_data == True,
        )
    else:
        statement = select(SourceFile).filter(
            SourceFile.account_unique_id == account_unique_id,
            SourceFile.included_in_source_data == True,
            SourceFile.already_processed_to_source_data == False
        )
        
    result = await session.execute(statement)
    documents_from_db = result.scalars().all()

    documents = []
    
    for db_file in documents_from_db:
        # Construct S3 key (path in S3) using account_unique_id and file name
        s3_key = f"{account_unique_id}/{db_file.file_name}"

        try:
            
            s3_object = s3.get_object(Bucket=BUCKET_NAME, Key=s3_key)

            file_content = s3_object['Body'].read()

            # Process file content based on its extension
            file_extension = os.path.splitext(db_file.file_name)[1].lower()
            content = await read_file_from_s3(file_content, file_extension)

            # Append the document with metadata
            documents.append(Document(
                page_content=content, 
                metadata={"file_name": db_file.file_name, "source": s3_key}
            ))

            # Mark file as processed in the database
            db_file.already_processed_to_source_data = True
            await session.commit()

        except ClientError as e:
            print(f"Failed to fetch file {db_file.file_name} from S3: {e}")
            continue
        except Exception as e:
            print(f"Failed to process file {db_file.file_name}: {e}")
            continue

    print(f"Loaded {len(documents)} documents from S3 based on DB query.")
    return documents


async def read_file_from_s3(file_content: bytes, file_extension: str) -> str:
    """
    Helper function to handle reading files from S3 based on the file extension.
    """
    print(f"Processing file with extension: {file_extension}")
    
    # Handling .txt files
    if file_extension == ".txt":
        try:
            return file_content.decode('utf-8')
        except Exception as e:
            print(f"Error decoding .txt file: {e}")
            raise

    # Handling .docx files
    elif file_extension == ".docx":
        try:
            # Convert the byte stream to an in-memory file-like object
            doc = DocxDocument(io.BytesIO(file_content))
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"Error processing .docx file: {e}")
            raise

    # Handling .doc files
    elif file_extension == ".doc":
        try:
            # Save the file content to an in-memory file
            temp_file = io.BytesIO(file_content)
            temp_file.name = "temp.doc"  # Required for pypandoc
            return pypandoc.convert_file(temp_file.name, 'plain')
        except Exception as e:
            print(f"Failed to convert .doc file: {e}")
            raise

    # Handling .md (markdown) files
    elif file_extension == ".md":
        try:
            markdown_content = file_content.decode('utf-8')
            return markdown.markdown(markdown_content)
        except Exception as e:
            print(f"Error decoding markdown file: {e}")
            raise

    # Handling .pdf files
    elif file_extension == ".pdf":
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            return "\n".join([page.extract_text() for page in pdf_reader.pages])
        except Exception as e:
            print(f"Error processing .pdf file: {e}")
            raise

    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


async def split_text(documents: list[Document]):
    """
    Split the text into chunks.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=500,
        length_function=len,
        add_start_index=True,
    )
    chunks = text_splitter.split_documents(documents)
    print(f"Split {len(documents)} documents into {len(chunks)} chunks.")
    return chunks


async def batch(iterable, n=BATCH_SIZE):
    """
    Helper function to split a list into batches of size n.
    """
    print(f"Batching iterable with size {len(iterable)} and batch size {n}.")
    l = len(iterable)
    for ndx in range(0, l, n):
        yield iterable[ndx:min(ndx + n, l)]
        await asyncio.sleep(0)
        
  