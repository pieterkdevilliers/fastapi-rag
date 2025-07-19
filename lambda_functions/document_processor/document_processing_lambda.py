# THIS MUST BE THE VERY FIRST THING IN THE FILE, BEFORE ANY OTHER IMPORTS.
__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')

import os
import io
import uuid
import gc
import boto3
import openai
import chromadb

# Import all necessary parsing and langchain libraries
from docx import Document as DocxDocument
import PyPDF2
import pypandoc
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_openai import OpenAIEmbeddings
from chromadb.api.types import EmbeddingFunction
from typing import Optional
import pandas as pd


# --- Configuration (Loaded from Lambda Environment Variables) ---
# The chromadb client will now read ALL the CHROMA_* variables automatically.
openai.api_key = os.environ['OPENAI_API_KEY']
BUCKET_NAME = os.environ['AWS_STORAGE_BUCKET_NAME']

# --- Global Clients (Initialized once per Lambda container start) ---
s3_client = boto3.client('s3')
textract_client = boto3.client('textract')

CHROMA_SERVER_AUTHN_CREDENTIALS = os.environ['CHROMA_SERVER_AUTHN_CREDENTIALS']
chroma_headers = {'X-Chroma-Token': CHROMA_SERVER_AUTHN_CREDENTIALS}

class ChromaEmbeddingFunction(EmbeddingFunction):
    """A wrapper for the LangChain OpenAIEmbeddings to be used by ChromaDB."""
    def __init__(self):
        self.embedding_function = OpenAIEmbeddings()

    def __call__(self, input: list[str]) -> list[list[float]]:
        return self.embedding_function.embed_documents(input)

# === THE LAMBDA HANDLER - MAIN ENTRY POINT ===

def handler(event, context):
    s3_bucket = event['s3_bucket']
    s3_key = event['s3_key']
    account_unique_id = event.get('account_unique_id', s3_key.split('/')[0])
    print(f"Starting processing for s3://{s3_bucket}/{s3_key}")
    try:
        file_content, file_extension = download_from_s3(s3_bucket, s3_key)

        # --- NEW LOGIC BRANCH FOR EXCEL FILES ---
        if file_extension in ['.xls', '.xlsx']:
            # For Excel, we parse directly into final chunks, skipping the split_text step.
            chunks = parse_excel_to_chunks(file_content, s3_key)

        else:
            text = parse_file_content(file_content, file_extension, s3_key)
            
            if text:
                source_document = Document(page_content=text, metadata={"source": s3_key})
                chunks = split_text([source_document])
            else:
                print(f"Parsing returned no text for {s3_key}. No chunks generated.")

        if chunks:
            save_chunks_to_chroma(chunks, account_unique_id)
        else:
            print("No chunks were generated. Nothing to save.")
        return {"statusCode": 200, "body": "File processed successfully."}
    except Exception as e:
        print(f"FATAL ERROR processing {s3_key}: {e}")
        raise e

# === HELPER FUNCTIONS ===

def save_chunks_to_chroma(chunks: list[Document], account_unique_id: str):
    """Connects to remote ChromaDB and saves chunks."""
    CHROMA_ENDPOINT = os.environ['CHROMA_ENDPOINT']
    print(f"Connecting to ChromaDB at {CHROMA_ENDPOINT}...")
    chroma_client = chromadb.HttpClient(
        host=CHROMA_ENDPOINT,
        headers=chroma_headers
    )

    print(f"Successfully connected to ChromaDB.")
    collection_name = f"collection-{account_unique_id}"
    embedding_function = ChromaEmbeddingFunction()
    collection = chroma_client.get_or_create_collection(
        name=collection_name,
        embedding_function=embedding_function
    )
    print(f"Using Chroma collection: {collection.name} with ID: {collection.id}")
    num_chunks = len(chunks)
    if num_chunks == 0:
        return
    collection.add(
        ids=[str(uuid.uuid4()) for _ in range(num_chunks)],
        documents=[chunk.page_content for chunk in chunks],
        metadatas=[chunk.metadata for chunk in chunks]
    )
    print(f"Successfully added {num_chunks} chunks to Chroma collection.")
    gc.collect()


def download_from_s3(bucket, key):
    print(f"Downloading {key} from bucket {bucket}...")
    s3_object = s3_client.get_object(Bucket=bucket, Key=key)
    file_content = s3_object['Body'].read()
    file_extension = os.path.splitext(key)[1].lower()
    return file_content, file_extension


def parse_file_content(file_content: bytes, file_extension: str, s3_key: str) -> Optional[str]:
    """
    Parses the file content based on its extension.
    Supports PDF, DOCX, TXT, MD, and DOC formats.
    """
    print(f"Parsing {s3_key} with extension: {file_extension}")
    if file_extension == ".pdf":
        return hybrid_parse_pdf(file_content, s3_key)
    elif file_extension == ".docx":
        doc = DocxDocument(io.BytesIO(file_content))
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_extension == ".txt":
        return file_content.decode('utf-8')
    elif file_extension == ".md":
        return markdown.markdown(file_content.decode('utf-8'))
    elif file_extension == ".doc":
        temp_file_path = f"/tmp/{uuid.uuid4()}.doc"
        with open(temp_file_path, "wb") as f:
            f.write(file_content)
        return pypandoc.convert_file(temp_file_path, 'plain')
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")
    

# NEW helper function specifically for PDFs
def hybrid_parse_pdf(file_content: bytes, s3_key: str) -> Optional[str]:
    """
    Parses a PDF. First tries a fast text extraction. If that fails or returns
    no text, it falls back to OCR with AWS Textract.
    """
    text = ""
    
    # --- STAGE 1: Try the fast, standard method first ---
    print(f"[{s3_key}] Attempting standard text extraction...")
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        for page in pdf_reader.pages:
            text += page.extract_text() or "" # Add 'or ""' to handle None return
    except Exception as e:
        print(f"[{s3_key}] Standard PDF parsing failed: {e}. Will proceed with OCR.")
        text = "" # Ensure text is empty if PyPDF2 fails

    # --- STAGE 2: If no text, use OCR as a fallback ---
    # We check if the stripped text is empty. A scanned PDF might return just whitespace.
    if not text.strip():
        print(f"[{s3_key}] No text found with standard method. Falling back to AWS Textract for OCR.")
        try:
            # For synchronous operations on single/multi-page PDFs
            response = textract_client.detect_document_text(
                Document={'Bytes': file_content}
            )
            # Reconstruct the text from Textract's response
            text = ""
            for item in response["Blocks"]:
                if item["BlockType"] == "LINE":
                    text += item["Text"] + "\n"
        except Exception as e:
            print(f"[{s3_key}] FATAL: AWS Textract OCR failed: {e}")
            return None # OCR failed, so we can't process this file
            
    print(f"[{s3_key}] Successfully parsed PDF. Extracted ~{len(text)} characters.")
    return text


def split_text(documents: list[Document]):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        add_start_index=True,
    )
    chunks = text_splitter.split_documents(documents)
    print(f"Split document into {len(chunks)} chunks.")
    return chunks


def parse_excel_to_chunks(file_content: bytes, s3_key: str) -> list[Document]:
    """
    Parses an Excel file and converts each row into a LangChain Document object.
    Each row becomes a separate chunk with its own metadata.

    :param file_content: The byte content of the .xls or .xlsx file.
    :param s3_key: The S3 key of the source file for metadata.
    :return: A list of LangChain Document objects.
    """
    print(f"Parsing Excel file {s3_key} with pandas...")
    all_chunks = []
    
    try:
        # Use ExcelFile to be able to access all sheets
        xls = pd.ExcelFile(io.BytesIO(file_content))
    except Exception as e:
        print(f"Pandas could not read the Excel file {s3_key}. It might be corrupt or password-protected. Error: {e}")
        return []

    for sheet_name in xls.sheet_names:
        print(f"Processing sheet: '{sheet_name}'")
        df = pd.read_excel(xls, sheet_name=sheet_name)
        
        # Filter out completely empty rows
        df.dropna(how='all', inplace=True)
        if df.empty:
            continue

        # Iterate over each row in the DataFrame
        for index, row in df.iterrows():
            # Convert the row to a descriptive string format
            # e.g., "Column1: Value1, Column2: Value2, ..."
            row_text_parts = []
            for col, val in row.items():
                if pd.notna(val) and str(val).strip(): # Ensure value is not null/empty
                    row_text_parts.append(f"{col}: {val}")
            
            row_text = ", ".join(row_text_parts)

            # Skip if the row ended up being empty after processing
            if not row_text:
                continue

            # Create rich metadata for each row
            row_metadata = {
                "source": s3_key,
                "sheet_name": sheet_name,
                # Add 2 to the index: +1 because index is 0-based, +1 for the header row
                "row_number": index + 2 
            }
            
            # Create a LangChain Document for this single row
            chunk = Document(page_content=row_text, metadata=row_metadata)
            all_chunks.append(chunk)

    print(f"Generated {len(all_chunks)} chunks directly from Excel file {s3_key}.")
    return all_chunks